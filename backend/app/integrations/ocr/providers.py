from __future__ import annotations

import logging
import re
from io import BytesIO
from abc import ABC, abstractmethod
from typing import Any, Optional

import httpx

logger = logging.getLogger(__name__)


class BaseOCRProvider(ABC):
    @abstractmethod
    async def parse_document(
        self,
        file_name: str,
        content_bytes: bytes,
        document_type: str = "resume",
        raw_text: Optional[str] = None,
    ) -> dict[str, Any]:
        raise NotImplementedError


def _extract_keywords(text: str, candidates: list[str]) -> list[str]:
    normalized_text = _normalize_for_match(text)
    return [
        item for item in candidates
        if _normalize_for_match(item) in normalized_text
    ]


def _normalize_ocr_text(text: str) -> str:
    # Some PDF text layers are extracted as "P y t h o n" or "数 据 分 析".
    # Collapse spaces inside continuous English/digit/Chinese tokens before parsing.
    text = re.sub(r"(?<=[A-Za-z0-9])[ \t]+(?=[A-Za-z0-9])", "", text)
    text = re.sub(r"(?<=[\u4e00-\u9fff])[ \t]+(?=[\u4e00-\u9fff])", "", text)
    return text


def _normalize_for_match(text: str) -> str:
    return re.sub(r"\s+", "", _normalize_ocr_text(text).lower())


def _extract_gpa(text: str) -> Optional[float]:
    match = re.search(r"(GPA|绩点)[:： ]*([0-9]\.?[0-9]*)", text, flags=re.IGNORECASE)
    return float(match.group(2)) if match else None


def _extract_major(text: str) -> str:
    explicit = re.search(r"(专业|Major)\s*[:：]\s*([^\n|｜，,；;]+)", text, flags=re.IGNORECASE)
    if explicit:
        return explicit.group(2).strip()

    known_majors = [
        "计算机科学与技术",
        "数据科学与大数据技术",
        "软件工程",
        "人工智能",
        "网络工程",
        "信息管理与信息系统",
        "数字媒体技术",
    ]
    for major in known_majors:
        if re.search(rf"{re.escape(major)}\s*专业", text):
            return major
    for major in known_majors:
        if major in text:
            return major

    degree_line = re.search(r"([^\n|｜]{2,30})\s*[|｜]\s*(本科|硕士|博士|专科)", text)
    if degree_line:
        candidate = degree_line.group(1).strip()
        if not any(token in candidate for token in ["GPA", "电话", "邮箱", "城市", "薪资"]):
            return candidate
    return ""


def _extract_name(text: str) -> str:
    explicit = re.search(r"姓名[:： ]*([^\n]+)", text)
    if explicit:
        value = re.split(r"意向岗位|电话|手机|邮箱|性别|出生|求职意向", explicit.group(1).strip())[0]
        return value.strip(" ：:")
    for line in text.splitlines():
        candidate = line.strip()
        if re.fullmatch(r"[\u4e00-\u9fa5·]{2,8}", candidate):
            return candidate
    return "未知学生"


def _extract_target_job(text: str) -> str:
    match = re.search(r"意向岗位[:： ]*([^\n]+)", text)
    if not match:
        return ""
    value = re.split(r"意向城市|期望薪资|求职类型|比赛经历|项目经历|教育背景|技能|证书|[，,；;]", match.group(1).strip())[0]
    return value.strip(" ：:")


def _extract_text_from_office_file(file_name: str, content_bytes: bytes) -> str:
    lowered = file_name.lower()
    if lowered.endswith(".pdf"):
        try:
            from pypdf import PdfReader

            reader = PdfReader(BytesIO(content_bytes))
            return "\n".join((page.extract_text() or "") for page in reader.pages).strip()
        except Exception as exc:
            logger.warning("Failed to extract PDF text for %s: %s", file_name, exc)
            return ""
    if lowered.endswith(".docx"):
        try:
            from docx import Document

            document = Document(BytesIO(content_bytes))
            lines = [paragraph.text for paragraph in document.paragraphs if paragraph.text.strip()]
            for table in document.tables:
                for row in table.rows:
                    row_text = " ".join(cell.text.strip() for cell in row.cells if cell.text.strip())
                    if row_text:
                        lines.append(row_text)
            return "\n".join(lines).strip()
        except Exception as exc:
            logger.warning("Failed to extract DOCX text for %s: %s", file_name, exc)
            return ""
    return ""


def _empty_ocr_result(document_type: str, message: str) -> dict[str, Any]:
    return {
        "raw_text": "",
        "layout_blocks": [],
        "structured_json": {
            "document_type": document_type,
            "name": "未知学生",
            "major": "",
            "skills": [],
            "certificates": [],
            "projects": [],
            "internships": [],
            "gpa": None,
            "ocr_warning": message,
        },
    }


class MockOCRProvider(BaseOCRProvider):
    SKILL_CANDIDATES = [
        "Python",
        "JavaScript",
        "TypeScript",
        "React",
        "Next.js",
        "FastAPI",
        "PostgreSQL",
        "Redis",
        "SQL",
        "MySQL",
        "Java",
        "Go",
        "C语言",
        "数据结构",
        "数据分析",
        "数据清洗",
        "大数据分析",
        "Docker",
        "Linux",
        "Figma",
        "机器学习",
        "深度学习",
        "PyTorch",
        "Excel",
        "ECharts",
        "Spring Boot",
        "数据可视化",
    ]
    CERTIFICATE_CANDIDATES = [
        "英语四级",
        "英语六级",
        "软件设计师",
        "计算机二级",
        "数据分析师证书",
        "产品经理证书",
        "人工智能工程师",
    ]

    async def parse_document(
        self,
        file_name: str,
        content_bytes: bytes,
        document_type: str = "resume",
        raw_text: Optional[str] = None,
    ) -> dict[str, Any]:
        text = raw_text or _extract_text_from_office_file(file_name, content_bytes)
        if not text:
            lowered = file_name.lower()
            if lowered.endswith((".pdf", ".doc", ".docx", ".png", ".jpg", ".jpeg")):
                return _empty_ocr_result(
                    document_type,
                    "未能从文件中提取可用文字，请使用真实 OCR 或上传可复制文本的 PDF/DOCX。",
                )
            text = content_bytes.decode("utf-8", errors="ignore")
        if not text.strip():
            return _empty_ocr_result(
                document_type,
                "未能从文件中提取可用文字。",
            )
        text = _normalize_ocr_text(text)
        skills = _extract_keywords(text, self.SKILL_CANDIDATES)
        certificates = _extract_keywords(text, self.CERTIFICATE_CANDIDATES)
        projects = re.findall(r"(项目|Project)[:： ]*(.+)", text)
        internships = re.findall(r"(实习|Internship)[:： ]*(.+)", text)
        structured = {
            "document_type": document_type,
            "name": _extract_name(text),
            "major": _extract_major(text),
            "target_job": _extract_target_job(text),
            "skills": skills,
            "certificates": certificates,
            "projects": [item[1].strip() for item in projects],
            "internships": [item[1].strip() for item in internships],
            "gpa": _extract_gpa(text),
        }
        layout_blocks = [
            {"section": "header", "text": structured["name"]},
            {"section": "skills", "text": ", ".join(skills)},
            {"section": "certificates", "text": ", ".join(certificates)},
            {"section": "experience", "text": "; ".join(structured["internships"] + structured["projects"])},
        ]
        return {
            "raw_text": text,
            "layout_blocks": layout_blocks,
            "structured_json": structured,
        }


class PaddleOCRProvider(BaseOCRProvider):
    def __init__(self, service_url: str, api_key: str = "") -> None:
        self.service_url = service_url.rstrip("/")
        self.api_key = api_key

    async def parse_document(
        self,
        file_name: str,
        content_bytes: bytes,
        document_type: str = "resume",
        raw_text: Optional[str] = None,
    ) -> dict[str, Any]:
        local_text = raw_text or _extract_text_from_office_file(file_name, content_bytes)
        if local_text:
            logger.info("Using local text extraction before PaddleOCR: file=%s", file_name)
            return await MockOCRProvider().parse_document(file_name, content_bytes, document_type, raw_text=local_text)

        headers = {}
        if self.api_key:
            headers["Authorization"] = f"Bearer {self.api_key}"
        try:
            async with httpx.AsyncClient(timeout=60) as client:
                logger.info("PaddleOCR request: POST %s file=%s", self.service_url, file_name)
                response = await client.post(
                    f"{self.service_url}",
                    files={"file": (file_name, content_bytes)},
                    data={"document_type": document_type, "raw_text": raw_text or ""},
                    headers=headers,
                )
                response.raise_for_status()
                result = response.json()
                logger.info("PaddleOCR success: file=%s status=%d", file_name, response.status_code)
                return result
        except httpx.HTTPStatusError as e:
            logger.error(
                "PaddleOCR HTTP error: file=%s status=%d body=%s",
                file_name, e.response.status_code, e.response.text[:500],
            )
            raise ValueError(f"PaddleOCR API returned {e.response.status_code}: {e.response.text[:200]}") from e
        except httpx.RequestError as e:
            logger.error("PaddleOCR connection error: file=%s url=%s err=%s", file_name, self.service_url, e)
            raise ValueError(f"PaddleOCR connection failed to {self.service_url}: {e}") from e
